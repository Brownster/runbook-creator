from flask import Flask, request, render_template, send_from_directory, redirect, url_for
import yaml
from docx import Document
from docx.shared import Pt
import os
import tempfile
import shutil

app = Flask(__name__)
UPLOAD_FOLDER = tempfile.mkdtemp()  # Uses temporary directory for file uploads and generated docs
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def add_run_with_font(paragraph, text, size):
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    return run

from docx.shared import Pt

def add_heading(doc, text, size, level):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.size = Pt(size)
    return heading

def add_paragraph(doc, title, content, title_size, content_size):
    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.size = Pt(title_size)
    para = doc.add_paragraph(content)
    run = para.add_run(content)
    run.font.size = Pt(content_size)
    return para

def create_runbook_file(file_path, original_filename):
    with open(file_path, 'r') as file:
        data = yaml.safe_load(file)
    
    if data is None or 'groups' not in data:
        raise ValueError("Invalid or empty YAML content")
    
    generated_filenames = []
    for group in data['groups']:
        for rule in group['rules']:
            alert_name = rule['alert']
            expr = rule['expr']
            description = rule['annotations']['description']
            severity = rule['labels']['severity']

            doc_filename = f"{original_filename}_{alert_name.replace(' ', '_')}.docx"
            doc_path = os.path.join(UPLOAD_FOLDER, doc_filename)
            doc = Document()

            # Correctly applying styles
            add_heading(doc, f"SM3 Alert RunBook – {group['name']} – {alert_name}", 14, level=1)
            add_paragraph(doc, 'Alert Name: ', alert_name, 12, 12)
            add_paragraph(doc, 'Alert Expression: ', expr, 12, 12)
            add_paragraph(doc, 'Category: ', group['name'], 12, 12)
            add_paragraph(doc, 'Description: ', description, 12, 12)
            add_paragraph(doc, 'Severity: ', severity, 12, 12)

            doc.save(doc_path)
            generated_filenames.append(doc_filename)
    
    return generated_filenames

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        if file:
            filename = file.filename
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(file_path)
            try:
                generated_filenames = create_runbook_file(file_path, filename)
                # Redirect to the download page for the first generated file
                return redirect(url_for('download_file', filename=generated_filenames[0]))
            except ValueError as e:
                return str(e), 400
    return render_template('upload.html')

@app.route('/downloads/<filename>')
def download_file(filename):
    try:
        # Correct usage without keyword arguments
        return send_from_directory(app.config['UPLOAD_FOLDER'], filename)
    except Exception as e:
        return str(e), 500

@app.route('/cleanup', methods=['POST'])
def cleanup():
    shutil.rmtree(UPLOAD_FOLDER)
    os.makedirs(UPLOAD_FOLDER)
    return redirect(url_for('upload_file'))

if __name__ == "__main__":
    app.run(debug=True)
